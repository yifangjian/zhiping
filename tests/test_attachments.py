"""附件處理(app/services/attachments.py、chat.build_user_content)。

核心原則:**不管他傳什麼,都不能沉默**。看不懂就誠實說看不懂,
但一定要有回應——在弱網路環境下,沒有回應會被解讀成「訊息沒送出去」。

測試資料全為虛構。
"""

import io

import pytest

from app.services import attachments
from app.services.chat import build_user_content
from app.services.location import from_location_message
from tests.fakes import FakeLine, make_message, make_runtime


def attachment(kind, **extra):
    message = make_message(text="")
    message["kind"] = kind
    message["message_id"] = "msg-1"
    message.update(extra)
    return message




def make_docx(paragraphs) -> bytes:
    import docx

    document = docx.Document()
    for line in paragraphs:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- 文件解析 ---


def test_讀得出_word_內容():
    data = make_docx(["排氣閥拆裝實習報告", "第一步:先洩壓"])

    text, note = attachments.extract_document_text(data, "作業.docx")

    assert "排氣閥拆裝實習報告" in text
    assert "先洩壓" in text
    assert note == ""


def test_word_表格內容也要抓():
    """規格數據常常在表格裡。"""
    import docx

    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "額定轉速"
    table.rows[0].cells[1].text = "750 rpm"
    buffer = io.BytesIO()
    document.save(buffer)

    text, _ = attachments.extract_document_text(buffer.getvalue(), "規格.docx")

    assert "額定轉速" in text and "750 rpm" in text


def test_太長的檔案會截斷並告知():
    data = make_docx(["很長的內容" * 2000])

    text, note = attachments.extract_document_text(data, "手冊.docx")

    assert len(text) <= attachments.MAX_EXTRACTED_CHARS
    assert "先看前面一部分" in note


def test_舊版_doc_誠實說讀不了():
    text, note = attachments.extract_document_text(b"\xd0\xcf\x11\xe0", "舊作業.doc")

    assert text is None
    assert ".docx" in note  # 要告訴他怎麼做才行


def test_不支援的副檔名():
    text, note = attachments.extract_document_text(b"xxxx", "圖檔.psd")

    assert text is None
    assert "讀不了" in note


def test_壞掉的_pdf_不會爆炸():
    text, note = attachments.extract_document_text(b"not a pdf at all", "壞的.pdf")

    assert text is None
    assert note  # 一定要有話說


def test_空內容的檔案建議改拍照():
    """掃描的手冊很常見:整頁都是圖,沒有文字層。"""
    data = make_docx([""])

    text, note = attachments.extract_document_text(data, "掃描.docx")

    assert text is None
    assert "拍" in note


# --- 對話紀錄裡的樣子 ---


def test_非文字訊息在紀錄裡有可讀的樣子():
    assert attachments.describe(attachment("image")) == "[圖片]"
    assert attachments.describe(attachment("sticker")) == "[貼圖]"
    assert (
        attachments.describe(attachment("file", file_name="手冊.pdf"))
        == "[檔案:手冊.pdf]"
    )
    assert (
        attachments.describe(attachment("location", title="釜山港"))
        == "[位置:釜山港]"
    )
    assert attachments.describe(make_message(text="嗨")) == "嗨"


# --- 組成模型輸入 ---


class ContentLine(FakeLine):
    """會回傳附件內容的假 LINE 用戶端。"""

    def __init__(self, content=None):
        super().__init__()
        self.content = content

    async def get_content(self, message_id):
        return self.content


async def test_圖片會以多模態格式交給模型():
    runtime = make_runtime(line=ContentLine(content=b"\xff\xd8\xff fake jpeg"))

    content = await build_user_content(runtime, [attachment("image")])

    assert isinstance(content, list)
    assert any(part["type"] == "input_image" for part in content)
    assert content[-1]["image_url"].startswith("data:image/jpeg;base64,")


async def test_圖片可以和文字一起送():
    runtime = make_runtime(line=ContentLine(content=b"fake"))
    batch = [make_message(text="這頁在講什麼"), attachment("image")]

    content = await build_user_content(runtime, batch)

    texts = [p["text"] for p in content if p["type"] == "input_text"]
    assert "這頁在講什麼" in texts[0]
    assert any(p["type"] == "input_image" for p in content)


async def test_檔案內容會變成文字附註():
    data = make_docx(["掃氣箱著火的處理程序"])
    runtime = make_runtime(line=ContentLine(content=data))

    content = await build_user_content(
        runtime, [attachment("file", file_name="程序.docx")]
    )

    assert isinstance(content, str)
    assert "程序.docx" in content
    assert "掃氣箱著火的處理程序" in content


async def test_下載失敗也要有話說():
    runtime = make_runtime(line=ContentLine(content=None))

    content = await build_user_content(runtime, [attachment("file", file_name="大檔.pdf")])

    assert "下載失敗" in content or "太大" in content


@pytest.mark.parametrize(
    "kind,keyword",
    [("sticker", "貼圖"), ("audio", "語音"), ("video", "影片")],
)
async def test_處理不了的型別也不沉默(kind, keyword):
    runtime = make_runtime(line=ContentLine())

    content = await build_user_content(runtime, [attachment(kind)])

    assert keyword in content


async def test_位置訊息會寫進提示():
    runtime = make_runtime(line=ContentLine())

    content = await build_user_content(
        runtime, [attachment("location", title="釜山港", address="Busan")]
    )

    assert "釜山港" in content


# --- 位置 → 時區 ---


def test_位置訊息用地址判斷時區():
    tz = from_location_message({"title": "釜山港", "address": "Busan", "longitude": 129.0})

    assert tz == "Asia/Seoul"


def test_公海上用經度估算時區():
    """認不出地名時還是要給個答案,誤差最多一小時,總比用台灣時間好。"""
    assert from_location_message({"address": "", "longitude": 120.0}) == "Etc/GMT-8"
    assert from_location_message({"address": "", "longitude": -60.0}) == "Etc/GMT+4"
    assert from_location_message({"address": "", "longitude": 0.0}) == "UTC"


def test_沒有座標也沒有地名就回_None():
    assert from_location_message({"address": "", "title": ""}) is None
