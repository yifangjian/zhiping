"""把使用者傳來的附件變成 LLM 讀得懂的東西。

**為什麼要做這個**:使用者的主要壓力是「要自己看說明書學」,而說明書多半是
全英文的 PDF 或掃描頁。讓他直接拍一頁、或把作業檔案丟過來問,比打字描述快得多——
尤其在機艙裡手很髒的時候。

處理方式依型別而異:

    圖片        原樣交給模型看(模型本身支援視覺)
    PDF / Word  抽出文字再交給模型
    其他        給一句誠實的回應,不要假裝看得到

**兩個上限**都是為了時間預算(見 docs/DESIGN.md 5.2):附件大小 10MB、
抽出的文字 6000 字。超過就截斷並告訴使用者,不要默默丟掉一半內容。

檔案內容全程只在記憶體裡,不落地——那是使用者的私人資料,少一個地方存
就少一個外洩的可能。
"""

import io
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# 抽出的文字上限。一份完整手冊塞進 prompt 會爆掉成本,也超出時間預算。
MAX_EXTRACTED_CHARS = 6000

# 圖片交給模型時用 data URL,不落地成檔案
SUPPORTED_IMAGE_TYPES = ("image/jpeg", "image/png", "image/gif", "image/webp")


def extract_document_text(data: bytes, file_name: str) -> Tuple[Optional[str], str]:
    """從 PDF 或 Word 抽出文字。

    回傳 (文字, 給使用者看的說明)。抽不出來時文字是 None,說明會講清楚原因——
    使用者需要知道是「檔案讀不到」還是「知平不想理他」。
    """
    lowered = (file_name or "").lower()

    if lowered.endswith(".pdf"):
        text = _read_pdf(data)
        kind = "PDF"
    elif lowered.endswith(".docx"):
        text = _read_docx(data)
        kind = "Word"
    elif lowered.endswith(".doc"):
        # 舊版 .doc 是二進位格式,純 Python 讀不了。誠實說,不要假裝
        return None, "這是舊版的 .doc,我讀不了。存成 PDF 或 .docx 再傳一次就可以。"
    elif lowered.endswith(".txt") or lowered.endswith(".md"):
        text = _read_plain(data)
        kind = "文字檔"
    else:
        return None, "這種檔案我讀不了,你可以把重點打出來問我。"

    if text is None:
        return None, "這份{}打不開,可能是檔案壞了或有加密。".format(kind)

    text = text.strip()
    if not text:
        # 掃描的手冊很常見:整頁都是圖,沒有文字層
        return None, (
            "這份{}裡面抓不到文字,可能是掃描的圖檔。"
            "你可以直接拍那一頁傳給我,我看得到圖。".format(kind)
        )

    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]
        return text, "檔案有點長,我先看前面一部分。"

    return text, ""


def _read_pdf(data: bytes) -> Optional[str]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
            # 邊讀邊檢查長度,不要為了一份 300 頁的手冊白跑完
            if sum(len(p) for p in parts) > MAX_EXTRACTED_CHARS:
                break
        return "\n".join(parts)
    except Exception:  # noqa: BLE001 — 壞檔案的形式太多種,一律當作讀不到
        logger.exception("讀取 PDF 失敗")
        return None


def _read_docx(data: bytes) -> Optional[str]:
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # 表格裡常有規格數據,一起抓
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:  # noqa: BLE001
        logger.exception("讀取 Word 失敗")
        return None


def _read_plain(data: bytes) -> Optional[str]:
    for encoding in ("utf-8", "big5", "utf-16"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    logger.warning("文字檔編碼認不出來")
    return None


def describe(message: dict) -> str:
    """一則訊息在對話紀錄裡的樣子。

    非文字訊息也要留下紀錄,否則之後撈短期脈絡時會看到莫名的空白,
    知平會不知道中間發生過什麼。
    """
    kind = message.get("kind", "text")
    if kind == "text":
        return message.get("text") or ""
    if kind == "image":
        return "[圖片]"
    if kind == "file":
        return "[檔案:{}]".format(message.get("file_name") or "未命名")
    if kind == "location":
        place = message.get("title") or message.get("address") or ""
        return "[位置:{}]".format(place) if place else "[位置]"
    if kind == "sticker":
        return "[貼圖]"
    if kind == "audio":
        return "[語音訊息]"
    if kind == "video":
        return "[影片]"
    return "[{}]".format(kind)


def to_data_url(data: bytes, content_type: str = "image/jpeg") -> str:
    """把圖片轉成 data URL 交給模型。

    用 data URL 而不是先上傳到某個空間再給網址:少一個外部相依、少一份
    使用者照片的副本存在別的地方,也省掉一次網路往返。
    """
    import base64

    encoded = base64.b64encode(data).decode("ascii")
    return "data:{};base64,{}".format(content_type, encoded)
